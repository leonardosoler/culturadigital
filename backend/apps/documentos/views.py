import io

from django.db.models import Q
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from docx import Document
from rest_framework import status, viewsets
from rest_framework.decorators import action
from rest_framework.response import Response

from apps.acompanhamento.models import AcompanhamentoEdital
from apps.ai_services import services as ai_services

from .models import DocumentoGerado, TemplateDocumento
from .serializers import DocumentoGeradoSerializer, GerarMinutaSerializer, TemplateDocumentoSerializer


class DocumentoGeradoViewSet(viewsets.ReadOnlyModelViewSet):
    """Minutas geradas por IA: listagem, geração e exportação em .docx."""

    serializer_class = DocumentoGeradoSerializer
    filterset_fields = ["acompanhamento", "tipo"]

    def get_queryset(self):
        return DocumentoGerado.objects.filter(
            acompanhamento__organizacao=self.request.user.organizacao
        ).select_related("acompanhamento")

    @action(detail=False, methods=["post"], url_path="gerar")
    def gerar(self, request):
        serializer = GerarMinutaSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        acompanhamento = get_object_or_404(
            AcompanhamentoEdital,
            pk=serializer.validated_data["acompanhamento"],
            organizacao=request.user.organizacao,
        )
        tipo = serializer.validated_data["tipo"]

        template = (
            TemplateDocumento.objects.filter(tipo=tipo)
            .filter(Q(organizacao=request.user.organizacao) | Q(organizacao__isnull=True))
            .order_by("-organizacao_id")
            .first()
        )
        instrucoes_extra = template.instrucoes_extra if template else ""

        try:
            conteudo = ai_services.gerar_minuta(
                tipo, acompanhamento.edital, acompanhamento.organizacao, instrucoes_extra=instrucoes_extra
            )
        except RuntimeError as exc:
            return Response({"detail": str(exc)}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
        except ValueError as exc:
            return Response({"detail": str(exc)}, status=status.HTTP_400_BAD_REQUEST)

        ultima_versao = (
            DocumentoGerado.objects.filter(acompanhamento=acompanhamento, tipo=tipo)
            .order_by("-versao")
            .values_list("versao", flat=True)
            .first()
        )
        documento = DocumentoGerado.objects.create(
            acompanhamento=acompanhamento,
            tipo=tipo,
            conteudo=conteudo,
            versao=(ultima_versao or 0) + 1,
        )
        return Response(DocumentoGeradoSerializer(documento).data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=["get"], url_path="exportar")
    def exportar(self, request, pk=None):
        documento = self.get_object()
        buffer = io.BytesIO()
        doc = Document()

        for linha in documento.conteudo.splitlines():
            texto = linha.strip()
            if texto.startswith("### "):
                doc.add_heading(texto[4:], level=3)
            elif texto.startswith("## "):
                doc.add_heading(texto[3:], level=2)
            elif texto.startswith("# "):
                doc.add_heading(texto[2:], level=1)
            elif texto.startswith(("- ", "* ")):
                doc.add_paragraph(texto[2:], style="List Bullet")
            elif texto:
                doc.add_paragraph(texto)
            else:
                doc.add_paragraph("")

        doc.save(buffer)
        buffer.seek(0)

        nome_arquivo = f"{documento.tipo}_{documento.acompanhamento_id}_v{documento.versao}.docx"
        return FileResponse(
            buffer,
            as_attachment=True,
            filename=nome_arquivo,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


class TemplateDocumentoViewSet(viewsets.ModelViewSet):
    """Templates/instruções extras por tipo de documento (globais ou da organização)."""

    serializer_class = TemplateDocumentoSerializer
    filterset_fields = ["tipo"]

    def get_queryset(self):
        return TemplateDocumento.objects.filter(
            Q(organizacao=self.request.user.organizacao) | Q(organizacao__isnull=True)
        )

    def perform_create(self, serializer):
        serializer.save(organizacao=self.request.user.organizacao)
